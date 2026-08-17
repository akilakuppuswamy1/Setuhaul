"""Build speaker script DOCX aligned to SetuHaul_FDE_Assignment_Presentation_updated.pptx."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "SetuHaul_FDE_Presentation_Speaker_Script.docx"

NAVY = RGBColor(0x0F, 0x27, 0x44)
GOLD = RGBColor(0xC4, 0x8A, 0x1A)
MUTED = RGBColor(0x4A, 0x5D, 0x73)
INK = RGBColor(0x10, 0x24, 0x3E)

SLIDES: list[tuple[str, str, str, str]] = [
    (
        "1",
        "Title — SETUHAUL",
        "About 45 seconds",
        """Good morning. This is SetuHaul: deterministic warehouse appointment coordination with conversational AI.

The problem we were given is a late or disrupted inbound driver talking to a warehouse that has scarce appointment capacity — limited hours, docks, and slot counts — and we still have to book correctly when two clients try to confirm at the same time.

The sentence I want you to remember for the whole talk is this: we use AI for language, and we use deterministic Python for operational decisions. The LLM is not an autonomous booking agent. It cannot decide feasibility, lock a dock, or confirm an appointment on its own.

Steps 1 through 9 are implemented, including hardening. The required classroom path is driver conversation plus concurrent scarce capacity. Step 9 is optional and read-only — it ranks trucks, it does not book.

We did not use LangChain, LangGraph, Redis, Kafka, or OR-Tools. The system of record is a frozen 16-table PostgreSQL schema.

I will walk the problem, the architecture, the confirmation path, the SH-1024 hero conversation, concurrency, and then the evidence.""",
    ),
    (
        "2",
        "Agenda",
        "About 30 seconds",
        """Here is the path through the deck.

First, the warehouse problem — why a chat message cannot be treated as a booking.

Second, architecture: the AI boundary, the layered stack, and the frozen schema.

Third, the authority path: feasibility, controlled proposals, and locked allocation.

Fourth, conversation: allowlisted tools over the SH-1024 hero flow.

Fifth, safety: concurrency, stale proposals, and adversarial hardening.

Sixth, evidence: tests, the operations console, how to demo, and honest limitations.

If time is short, the three slides that carry the design are the AI boundary, show-versus-propose-versus-confirm, and the lock order.""",
    ),
    (
        "3",
        "Context — The receiving dock cannot treat chat as a booking",
        "About 75 seconds",
        """Start from the operational world, not from the chatbot.

A receiving facility has scarce appointment capacity. Our classroom story is shipment SH-1024. Jane Rivera was supposed to hit a 6:30 PM window. She is two hours late, so the new ETA is around 8:30 PM. She also has an emergency and must leave by 9:30 PM. Any slot that finishes after that is operationally useless even if the dock is empty.

Docks and slots have finite capacity. Several trucks can be eligible for the last remaining window. An option that looked feasible when we showed it can be gone by the time the driver says “confirm.” Two clients can both see remaining capacity of one, and both try to confirm.

That is why language is not authority. A chat message is not a lock, not a hold, and not a booking. If we treat “book me 8:30” as a write, we will double-book.

SetuHaul’s job is to record facts, evaluate feasibility with explicit rules, propose without consuming capacity, and confirm only after revalidation plus a transactional allocation.""",
    ),
    (
        "4",
        "Failure mode — Naive “book the slot I asked for” is unsafe",
        "About 70 seconds",
        """Look at the left column. That is the naive path.

The driver says “book me 8:30 PM.” The chat treats the message as a booking. There is no re-check at commit time. There are no row locks on the slot or the dock. Two confirms both succeed. The last unit of capacity is double-booked. Showing an option is not holding capacity. Feasibility at display time is not feasibility at commit time.

Now the right column — what this repository actually does.

We record facts first: the new ETA, and leave-by as conversation context. We evaluate with Step 5 rules. We create a proposal that does not consume capacity. We require explicit confirm language. On accept we revalidate Step 5, take locks in Step 6, and re-check capacity. One winner gets confirmed. The loser gets HTTP 409 or a stale proposal. There is no silent retry that pretends success.

That contrast is the whole assignment.""",
    ),
    (
        "5",
        "Assignment — Classroom problem mapped onto this repository",
        "About 60 seconds",
        """This table is how we mapped the FDE challenge onto code. I am not inventing extra requirement IDs.

Driver conversation is Step 8, over frozen ChatThread and ChatMessage tables. Scarce capacity is slot capacity, dock availability, and facility rules. Feasibility is Step 5 FeasibilityEngine — pure Python rules, no SQL, no LLM. Controlled proposals are Step 7: Appointment rows with status requested, not a separate Proposal table. Concurrency is Step 6 AllocationService: shipment then slot then dock locks, plus a capacity re-check. Optional facility scheduling is Step 9: read-only ranking, and there is no schedule-confirm endpoint.

The required classroom path is Steps 5 through 8. Step 9 never books, locks, reserves, or confirms capacity. A recommended slot still has to walk the same accept path.""",
    ),
    (
        "6",
        "Core design — LLM for language. Deterministic Python for decisions.",
        "About 80 seconds",
        """This is the design principle. Please treat the two columns as a contract.

On the left — what AI does. It understands driver language: intent, clock times, “the second one.” It classifies intent and extracts conversational constraints: ETA, leave-by, earliest start. It keeps context in ChatMessage metadata JSON plus the thread foreign keys. It may choose only allowlisted tools. It explains deterministic rule_results in driver-facing English. It may request human escalation as a record or flag — it does not dispatch a human.

On the right — what AI does not do. It does not decide feasibility. It does not invent travel times. It is not the capacity authority. It does not choose or lock docks as a commit. It does not allocate or confirm independently. It does not bypass concurrency controls, does not write SQL, and does not override Step 5, 6, or 7.

If a reviewer asks “can the model just confirm?” the answer is no. Privileged confirmation is only accept_proposal going into Step 7 accept, Step 5 revalidation, and Step 6 allocation.""",
    ),
    (
        "7",
        "Solution — Required confirmation path (driver → booking)",
        "About 70 seconds",
        """Walk the top row left to right. This is the only path that creates a confirmed appointment from a driver conversation.

The driver talks in the React operations console. The console posts to FastAPI. Step 8 is the conversation layer: intent, tools, wording. Step 5 evaluates feasibility when we list options. Step 7 writes a proposal. Then — and this is the important second Step 5 — we revalidate at accept time, because the world may have changed. Only then does Step 6 allocate under locks and write the confirmed appointment.

The gold bar is the sentence I will repeat: privileged confirmation is Step 8 accept_proposal, then Step 7 accept, then Step 5 revalidation, then Step 6 allocation.

The bottom row is optional Step 9. Facility snapshot, scheduling service, Step 5 eligibility, ranking engine, a proposed schedule — and then a red box: no booking, no lock. Step 9 is advice. Booking still comes back to the top row.""",
    ),
    (
        "8",
        "Architecture — Layered stack — UI never calculates feasibility",
        "About 70 seconds",
        """Each band only talks downward for writes. The UI never bypasses the API, and it does not calculate feasibility, capacity, dock compatibility, or confirmation.

Clients: React 19 plus Vite. Driver Console, Shipments, Appointments, Facility Schedule, Demo Scenarios, and a Concurrency page that explains the race — it does not fire two live confirms.

HTTP API: FastAPI, Pydantic contracts, CORS, and GET /health returning service setuhaul so the console can reject a different process on the same port. We run on 8010, not 8000.

Conversation: intent, entities, thread context, allowlisted tools, formatter. Default provider is FakeLLM.

Orchestration services load facts from repositories. They never invent operational measurements.

Decision cores: FeasibilityEngine is pure rules. Allocation is a transactional service with locks — not a second pure engine. SchedulingEngine only ranks.

Persistence: PostgreSQL 16, SQLAlchemy 2, Alembic, frozen 16-table Step 2 schema.""",
    ),
    (
        "9",
        "Diagram — Implemented layered stack",
        "About 55 seconds",
        """This is the official system architecture diagram from docs, the layered-stack section.

Point at the five bands from top to bottom. Clients have no decision logic. The HTTP API is the contract. Orchestration services load facts. Decision cores are FeasibilityEngine, allocation locks, SchedulingEngine, and the language adapter. Persistence is repositories, SQLAlchemy models, and PostgreSQL.

Say this out loud: “If you remember one architecture picture, remember that writes go down this stack. The React console displays API results. It is not a second feasibility engine.”

LangChain and LangGraph are not in this picture because they are not in the repository.""",
    ),
    (
        "10",
        "Diagram — Runtime request path and confirmation sequence",
        "About 70 seconds",
        """This is the same architecture PNG, the runtime and confirmation section.

Two clients share the same engines: free text through the conversation agent, and structured REST from the other console pages. Both paths hit Step 5 for eligibility. Both paths hit Steps 6 and 7 to consume capacity. Step 9 only reads.

Then walk the sequence. The driver sends free text. FakeLLM or OpenRouter extracts intent. It does not decide feasibility. The agent calls get_available_options or evaluate_feasibility. Step 5 returns ordered rule_results. create_proposal writes Appointment status requested. Capacity is not consumed. When the driver confirms, accept_proposal re-runs Step 5. If the world changed: stale, HTTP 409. If still feasible: Step 6 takes slot and dock row locks and writes the confirmed appointment.

There is no AI booking pathway on this diagram.""",
    ),
    (
        "11",
        "AI boundary — Allowlisted tools are the only crossing",
        "About 65 seconds",
        """Three columns.

Language side may parse free text, ask when facts are missing, explain rule_results, and flag human escalation.

The middle is the wall. Allowlisted tools only. An unknown tool name is forbidden. Tool arguments use Pydantic extra equals forbid. UUIDs are validated. If we detect injection markers, we skip irreversible tools.

Authority side must not invent. FeasibilityEngine, ProposalService, AllocationService, SchedulingEngine, and the ETA and exception services own the facts and the commits.

The LLM never jumps the wall with raw SQL or an arbitrary function call. Every operational effect is a named tool that already exists in the backend.""",
    ),
    (
        "12",
        "Diagram — Language side vs authority side",
        "About 60 seconds",
        """This is the official AI responsibility-boundary diagram.

On the left, FakeLLM or OpenRouter plus intents.py. Their job is parse, clarify, explain, flag. They must not invent facts, write SQL, or call arbitrary functions.

In the middle, ToolName plus ToolExecutor. Read and evaluate tools are safe to call when the intent matches. Irreversible write tools — record_eta_update, create_driver_exception, create_proposal, reject_proposal, accept_proposal — are skipped if injection is detected.

On the right, the Python engines. Step 5 eligibility, Step 6 row locks, Step 7 requested rows, Step 9 ranking. LangChain and LangGraph are not used.

If someone asks who confirms scarce capacity: accept, then Step 5 revalidate, then Step 6. The LLM may not independently confirm.""",
    ),
    (
        "13",
        "Diagram — No AI booking pathway",
        "About 55 seconds",
        """Stay on the same diagram, lower section.

Four boxes. Show uses get_available_options and FeasibilityEngine. Options are not a hold. Propose uses create_proposal. Status is requested. Capacity is still not consumed. Confirm uses accept_proposal, revalidates, then Step 6 locks and confirms. Step 9 rank is read-only and never holds or books.

Then the ownership table: parse and clarify — yes. Explain wording — yes. Shipment facts, slot eligibility, capacity confirm, facility ranking — no, except through the allowlisted tool, and confirm is never independent.

Same engines on REST. Driver console is free text. Other pages call REST. They share ProposalService. That is why the UI cannot cheat.""",
    ),
    (
        "14",
        "Language adapters — Two providers. Neither can authorize confirmation.",
        "About 55 seconds",
        """FakeLLMProvider is the default. It is a deterministic parser in intents.py. No network. LLM_PROVIDER equals fake is the safe classroom setting. The live hero flow uses this. Confirm and reject flags come from the parser.

OpenRouterProvider is optional language understanding. It cannot independently authorize confirmation. If the model invents a write intent that the parser did not also see, we reject it. Confirm and reject still come from the parser, not from the model payload. Missing API key falls back to FakeLLM.

LangChain and LangGraph are not used. There is no AI booking pathway. OpenRouter cannot promote a non-write parser intent to ACCEPT_PROPOSAL.""",
    ),
    (
        "15",
        "System of record — Frozen Step 2 schema — 16 tables, zero added later",
        "About 60 seconds",
        """PostgreSQL is the system of record. We froze the schema at Step 2 and 2H. Steps 3 through 9 add zero tables and zero new Alembic upgrade operations.

Five clusters. Actors: carriers, drivers, vehicles, contacts. Move: shipments, eta_updates, driver_exceptions. Facility: facilities, docks, facility_rules, appointment_slots. Commitment: appointments and facility_checkins. Conversation: chat_threads, chat_messages, operational_messages.

UUID primary keys. created_at on every table. Proposals are not a new table — they reuse appointments with status requested and a STEP7_PROPOSAL marker in notes. Conversation reuses chat_threads and chat_messages. Latest ETA is derived from eta_updates history; it is never stored on Shipment.

Alembic revisions are the Step 2 create and the 2H indexes and constraints. alembic check confirms the freeze held.""",
    ),
    (
        "16",
        "Diagram — Entity-relationship model — 16 frozen tables",
        "About 65 seconds",
        """This is the official ER diagram.

Point at Shipment in the center. Carrier, driver, and vehicle hang off the left. ETA updates and driver exceptions hang off the move. Facility, docks, rules, and slots hang off the right. Appointment connects shipment, facility, slot, and dock — that one table is both proposals and bookings. ChatThread hangs off shipment and driver; ChatMessage carries metadata JSON for Step 8 context.

Solid lines are required foreign keys. Dashed lines are optional. We did not add a Proposal table later because the assignment froze the schema.

Say: “Feasibility and ranking only read these rows. Allocation locks slots and docks, then writes Appointment. Conversation does not get its own operational tables.”""",
    ),
    (
        "17",
        "Build — Steps 1–9 — all implemented, including hardening",
        "About 55 seconds",
        """Quick roadmap so the later deep-dives have a map.

Step 1 is a runnable FastAPI with config and health. Step 2 is the 16-table system of record. 2H is indexes and constraints. Step 3 is read-only business APIs with pagination. Step 4 writes ETA and exceptions as facts. Step 5 is the pure feasibility engine. Step 6 is concurrency-safe allocation; 6H hardens it. Step 7 is controlled proposals. Step 8 is conversational AI; 8H is adversarial and confirmation hardening. Step 9 is optional read-only ranking.

Gold cards are hardening or the optional extension. Everything on this slide is implemented, not a future design.""",
    ),
    (
        "18",
        "Step 4 — Record operational facts — delay is not a booking",
        "About 70 seconds",
        """A driver saying “I will be two hours late” is a fact, not a booking.

Top row: driver reports delay, POST to shipments id eta-updates, we insert an immutable eta_updates row, latest ETA is derived, Step 5 later reads that latest ETA.

Left box: ETA updates. POST writes history. GET latest-eta derives it. Conversation tool is record_eta_update. That tool does not evaluate feasibility and does not consume a slot.

Right box: driver exceptions. POST creates them. open or acknowledged exceptions block Step 5 under rule EXCP-001. Important classroom nuance: “I need to leave by 9:30 PM” is stored as leave_by_local in conversation metadata. It is not automatically a blocking exception type. PATCH acknowledges or resolves. Still not a booking.

If you only remember one Step 4 sentence: driver-reported delay is a fact.""",
    ),
    (
        "19",
        "Diagram — ETA and exception writes — facts, not bookings",
        "About 50 seconds",
        """Official eta_exception_flow diagram.

Walk the write path. The conversation or REST client inserts eta_updates. History is append-only. Shipment is not mutated with a latest-ETA column. Exceptions are their own rows with status.

Point at the rule that open and acknowledged exceptions make a combination not feasible. Resolved exceptions do not block.

This diagram is here so a reviewer cannot claim we “booked from the chat message.” We wrote a fact row, then later feasibility reads it.""",
    ),
    (
        "20",
        "Diagram — How facts reach feasibility",
        "About 45 seconds",
        """Same PNG, lower section.

Step 8 tools record facts. Step 5 reads latest ETA and active exceptions when it evaluates a shipment-slot-dock combination. Missing facts yield not_evaluable. We do not invent travel times to fill gaps.

Neither of these writes is a capacity commit. Capacity still waits for Step 6 after a Step 7 accept.""",
    ),
    (
        "21",
        "Step 5 — FeasibilityEngine — pure rules, no database, no LLM",
        "About 80 seconds",
        """POST shipments id feasibility goes to FeasibilityService, which loads facts and calls FeasibilityEngine. Outcomes are feasible, not_feasible, or not_evaluable, plus ordered rule_results.

Walk the rule families. SHIP 001 to 003: shipment active, not terminal, destination assigned. Carrier, driver, vehicle must be active; weight and volume when data is present. Facility must be active. Slot must exist, match facility, be open, and have remaining capacity. Dock presence, facility, availability, weight, reefer compatibility. Facility rules: max daily appointments, operating hours, dock compatibility for vehicle type and pallet limits. ETA-001 is blocking if latest ETA is outside the slot window. ETA-002 is a warning if outside hours. EXCP-001 blocks on open or acknowledged exceptions.

Capacity-consuming appointment statuses are confirmed and held only. requested proposals do not consume capacity. That is how we can show and propose without locking the last slot.

The engine has no SQL and no LLM. Step 6 must re-run Step 5 inside the lock. Step 6 is not a second feasibility engine.""",
    ),
    (
        "22",
        "Step 7 — Showing is not proposing. Proposing is not confirming.",
        "About 70 seconds",
        """Three columns. This is the product contract.

Show: get_available_options. Step 5 feasibility. Numbered options in chat. No database write. Capacity unchanged. Driver-facing: I found these feasible options, which would you prefer?

Propose: create_proposal. Appointment status requested, STEP7_PROPOSAL in notes. Thirty-minute application TTL from created_at — there is no expires_at column. Still not a hold. Driver-facing: I have created a proposal, say confirm if you want me to book it.

Confirm: accept_proposal. Revalidate Step 5. Step 6 row locks. A new confirmed appointment. The proposal row is reconciled. Driver-facing: the appointment is confirmed.

You cannot jump from show to confirmed in one tool. There is no conversation tool named allocate.""",
    ),
    (
        "23",
        "Diagram — Three states as implemented",
        "About 50 seconds",
        """Official Show-Propose-Confirm diagram, the three-state section.

Driver or REST asks to show options. Picking a number creates a proposal. Confirm goes through feasibility re-check and allocation to a confirmed appointment. REST equivalent is POST proposals id accept.

Say: “If I show you three slots in chat, I have not reserved any of them. If I create a proposal, I still have not reserved the slot. Only accept consumes capacity.”""",
    ),
    (
        "24",
        "Diagram — Conversation sequence — show / propose / status / confirm",
        "About 70 seconds",
        """Same PNG, conversation sequence.

Turn pattern of the hero flow. ASK_OPTIONS calls get_available_options, Step 5 returns ordered feasible options. PROPOSE_CHANGE creates the proposal and writes Appointment requested. GET_STATUS or “has it been confirmed?” is ASK_STATUS — get_proposal only, no write. ACCEPT_PROPOSAL calls accept_proposal, re-checks feasibility, allocates under locks, returns confirmed_appointment_id.

Stress ASK_STATUS. A status question must not be misclassified as confirm. That is Step 8H routing, and we test it.

The Driver Console is only displaying this. The agent is calling the same services REST would call.""",
    ),
    (
        "25",
        "Diagram — REST accept internals — same ProposalService",
        "About 65 seconds",
        """Same PNG, REST and accept internals.

POST shipments id proposals: FeasibilityService first. If not feasible, 400. Else create requested appointment, 201 status proposed.

GET proposals id resolves status, and may persist expired if the 30-minute TTL passed.

POST proposals id accept: re-evaluate, allocate. Slot gone or feasibility changed: 409. Success: 200 confirmed.

Inside accept: advisory lock on shipment_id. Already confirmed: idempotent recovery. Expired, rejected, or stale: 409. Proposed: revalidate. Then AllocationService.allocate. On success the proposal is cancelled in favor of a separate confirmed row. No PATCH status endpoint. No silent retry.""",
    ),
    (
        "26",
        "Step 7 lifecycle — Proposal states — no Proposal table",
        "About 60 seconds",
        """API states. Create, if Step 5 is feasible, lands on proposed — that is the requested row.

From proposed you can reject, expire after 30 minutes, go stale if the world changed on accept, or confirm through the in-call accepted transition.

accepted is never persisted and never returned by GET. One accept call either confirms, marks stale, or fails.

Capacity is consumed only on the confirmed appointment row written by Step 6. Terminal states have no outbound edges. We do not mutate status with PATCH. HTTP 409 when the world changed.""",
    ),
    (
        "27",
        "Diagram — Proposal API states and transitions",
        "About 50 seconds",
        """Official proposal state diagram, the transition graph.

Point at proposed in the middle. Rejected, expired, and stale are terminals with no capacity. Confirmed is a terminal with capacity on the other row. The transient accepted node is labeled in-call only.

This graph is encoded as _VALID_TRANSITIONS in ProposalService. The LLM cannot invent a new state.""",
    ),
    (
        "28",
        "Diagram — Persistence mapping and accept path",
        "About 55 seconds",
        """Same PNG, mapping and accept path.

proposed maps to Appointment status requested. stale maps to cancelled plus a reason in notes. The proposal row after a successful confirm is cancelled and linked to confirmed_appointment_id. The booking itself is a separate confirmed or held row.

Accept path: lock, already-confirmed check, must still resolve to proposed, re-run Step 5, allocate Step 6, then mark the proposal done.

REST and conversation share this service: create_proposal, get_proposal, accept_proposal, reject_proposal.""",
    ),
    (
        "29",
        "Step 6 — Concurrency-safe allocation — one winner for the last unit",
        "About 70 seconds",
        """Two requests both see remaining capacity of one. Show and propose do not lock.

Then one transaction. Lock order is always shipment, then slot, then dock, to prevent deadlocks. Inside: advisory lock on the shipment, FOR UPDATE on appointment_slots, FOR UPDATE on docks, re-run Step 5 inside the lock, count confirmed plus held versus slot.capacity, write confirmed, safe_commit.

Winner: confirmed appointment, capacity consumed. Loser: Step 5 or capacity re-check fails, HTTP 409, stale or conflict, no silent retry, no false success.

Conversation never calls allocate directly. Confirmation uses Step 7 accept, which wraps this service.""",
    ),
    (
        "30",
        "Diagram — Lock order — shipment, then slot, then dock",
        "About 55 seconds",
        """Official concurrency diagram, lock-order section.

Candidates are resolved with unlocked reads first so we do not hold locks while listing open slots. Then pg_advisory_xact_lock on the shipment key. SQLite tests use SELECT shipments FOR UPDATE instead. Then slot FOR UPDATE, then dock FOR UPDATE. Then feasibility inside the lock. COMMIT releases everything.

ALLOCATION_LOCK_ORDER is shipment, slot, dock. Slot is always locked before dock.

This is why allocation is a service with locks, not a second pure engine.""",
    ),
    (
        "31",
        "Diagram — Capacity-1 race — two shipments, one slot",
        "About 55 seconds",
        """Same PNG, the capacity-1 race.

AllocationService A and B both want the last unit. A locks the slot row first. PostgreSQL makes B wait. A re-checks, writes confirmed, commits. B’s lock is granted, B counts confirmed plus held, capacity is gone, ConflictError.

Winner is whoever locked the slot row first — not whoever showed the option first, and not whoever the LLM preferred.

We have real PostgreSQL tests for this: test_step6_concurrency and test_step7_concurrency, thread pool, no silent retry.""",
    ),
    (
        "32",
        "Diagram — Two requested proposals, one confirm",
        "About 50 seconds",
        """Same PNG, Step 7 wrapping Step 6.

Two clients can both have a requested proposal for the last slot because requested does not consume capacity. That is intentional. The race is at accept, not at create.

Both call accept. One takes the shipment advisory lock, revalidates, allocates, confirms. The other finds the slot full or the proposal stale, HTTP 409.

A second conversation after SH-1024 is already confirmed cannot silently confirm again.""",
    ),
    (
        "33",
        "Step 8 — Allowlisted conversation tools — no allocate tool",
        "About 70 seconds",
        """These are the only tool names. Unknown names are rejected. There is no eval, exec, or SQL from the model.

Reads: get_shipment_status, get_proposal. Fact writes: record_eta_update, create_driver_exception. Evaluate: evaluate_feasibility, get_available_options, and optional evaluate_facility_schedule which is read-only. Proposal writes: create_proposal, reject_proposal. Escalation is a flag: request_human_escalation.

The red card is the only capacity commit from conversation: accept_proposal. It goes ProposalService.accept, Step 5, Step 6.

There is no conversation tool named allocate. Direct POST shipments id allocate exists for structured clients. The driver chat does not use it.

get_available_options may return dock_id null until a proposal or allocation assigns a dock. That is a documented boundary, not a bug.""",
    ),
    (
        "34",
        "Hero flow — SH-1024 — Jane Rivera, Dallas DC",
        "About 90 seconds",
        """This is the verified live hero flow from scripts/e2e_hero_flow.py against seeded shipment SH-1024.

Turn 1: “I will be two hours late, I was supposed to reach by 6:30, I will reach around 8:30.” Intent UPDATE_ETA or REPORT_DELAY. Tool record_eta_update. Authority is Step 4 fact write. Database: new eta_updates row.

Turn 2: “I also have an emergency and need to leave by 9:30 PM.” We store leave_by_local in metadata. We do not treat that phrasing as a blocking exception. No proposal yet.

Turn 3: “My ETA is 8:30. What options do I have?” ASK_OPTIONS, get_available_options, Step 5 plus leave-by filter. Showing is not proposing.

Turn 4: “The second one works, but I need to leave by 9:30.” PROPOSE_CHANGE, create_proposal. requested row. Confirmation has not occurred.

Turn 5: “Has it been confirmed?” ASK_STATUS, get_proposal, read only.

Turn 6: “Confirm it.” ACCEPT_PROPOSAL, accept_proposal, Step 7 then 5 then 6, confirmed appointment.

If you demo live, this is the script. Use Demo Scenarios or type these messages in the Driver Console.""",
    ),
    (
        "35",
        "Diagram — Driver conversation sequence (hero flow)",
        "About 55 seconds",
        """Official driver conversation sequence.

Point at each turn’s intent and tool as you retell SH-1024 in one breath: delay is a fact, leave-by is context, options are a read, proposal is requested, status is a read, confirm is the only commit.

The model never writes SQL and never confirms on its own. Every arrow that touches the database goes through an existing service.""",
    ),
    (
        "36",
        "Diagram — One-turn internals through the agent",
        "About 55 seconds",
        """Same PNG, internals.

POST conversations thread messages. ConversationAgent. Provider returns understanding JSON. If shipment, option, or delay is missing: clarification, no tools. If injection plus irreversible: skip write tools. Plan tools, ToolExecutor. Unknown name: forbidden. Then existing services and engines. Formatter produces driver-facing text. Prompts are never stored on ChatMessage.

This is why we can swap FakeLLM for OpenRouter without creating an AI booking pathway. The guards are in code, not only in the system prompt.""",
    ),
    (
        "37",
        "Diagram — End-to-end journey — SH-1024 at Dallas DC",
        "About 55 seconds",
        """Official end-to-end driver journey.

Starting world already exists from seed_ops_demo.py: carrier SETU-DEMO, driver Jane Rivera, vehicle, Dallas Distribution Center, original 6:30 to 7:00 appointment, later open slots. We do not invent a shipment in the chat.

Then the journey: delay, leave-by, options, proposal, status, confirm. UI displays API results only. Facility check-in exists as a read and ranking input; the conversation does not check the driver in at the gate.

No driver authentication. Escalation is a flag, not a dispatcher.""",
    ),
    (
        "38",
        "Diagram — From proposal to confirmed appointment",
        "About 50 seconds",
        """Same journey PNG, the commit half.

After the proposal exists, status questions stay read-only. Confirm revalidates. If another truck took the slot, stale, 409. If still feasible, locks, confirmed row, proposal reconciled.

Call out the two-row design: the proposal row is cancelled; Step 6 writes a separate confirmed appointment. That is how we keep the frozen schema and still distinguish propose from book.

The e2e script confirms SH-1024. Re-seeding does not undo that, because unique codes skip insert. Reset is docker compose down -v, not hand-editing rows.""",
    ),
    (
        "39",
        "Step 9 optional — Read-only facility ranking — never a hold",
        "About 65 seconds",
        """Optional extension. POST facilities facility_id schedule evaluate.

Input: facility_id, optional window, shipment ids, evaluated_at. Load a snapshot: shipments, open slots, docks, check-ins, ETAs. Eligible combinations via Step 5. Deterministic ranking. Response includes proposed_assignments, unassigned_shipments, read_only true, commits_capacity false.

Ranking as coded: confirmed or held appointments are protected and not moved. Then earlier facility check-in, then en-route. Lower ETA lateness versus slot end — missing ETA is not fabricated, those rank after known ETAs. Lower early-wait, closer ETA-to-slot-start, then stable ids. Score 0 to 100 from evaluable ETA metrics; null if ETA missing. Frozen schema has no shipment priority and no expected_unload_minutes.

Caps: 50 shipments, 100 slots. No OR-Tools. No POST schedule confirm. A recommended slot still confirms through Step 7, 5, and 6.""",
    ),
    (
        "40",
        "Diagram — Step 9 scheduling architecture",
        "About 45 seconds",
        """Official scheduling architecture diagram.

Facility snapshot into SchedulingService into Step 5 eligibility into SchedulingEngine into a proposed schedule that does not write appointments.

Point at read_only true and commits_capacity false. Recalculation means call evaluate again after operational state changes. It is not a live hold that drifts.""",
    ),
    (
        "41",
        "Diagram — Ranking rules and bounds",
        "About 45 seconds",
        """Same PNG, ranking section.

Walk the tie-break order slowly. Protected confirmed and held. Check-in before en-route. Do not invent missing ETAs. Stable sort keys at the end so the ranking is deterministic and testable.

No OR-Tools, no MIP, no second allocation engine. That is out of scope, not a missing feature of Step 9.""",
    ),
    (
        "42",
        "Operations console — React displays API results",
        "About 50 seconds",
        """frontend is Vite, React 19, TypeScript, React Router.

Slash is Driver Console — free text, shows intent, tools, read-only status. Shipments and Appointments are catalogs. Facility Schedule calls evaluate and labels it read-only. Demo Scenarios has the hero messages. Concurrency explains the race and does not fire two live confirms.

Stale and 409 are surfaced; the UI does not retry them as success. Escalation shows a flag; it does not claim a human has acted.

Five Vitest tests passed. Production build passed. No local feasibility or confirmation math.""",
    ),
    (
        "43",
        "Hardening — Adversarial and confirmation controls in the code",
        "About 70 seconds",
        """These are verified in Step 8H and related tests, not slide-ware.

Allowlisted tools. Strict arguments, extra forbid, UUID validation. Thread create rejects a shipment assigned to another driver. Allocation, proposal, and Step 9 reject mismatched facility, slot, or dock. Prompt injection: marker detection, skip irreversible tools. OpenRouter cannot promote confirm on its own. No eval, exec, or dynamic SQL from AI. public_metadata strips api_key, authorization, prompts. Conversation errors are generic — no traceback or SQLAlchemy leak. Stale confirm is 409, no silent retry. Status questions classify as ASK_STATUS and never call accept_proposal.

Honest remaining limitation: classroom authentication. Endpoints do not require login or JWT. The console says “No authentication.” Driver identity is a request field, not a verified session. I will repeat that on the limitations slide so it is not hidden.""",
    ),
    (
        "44",
        "Evidence — What was verified for this submission",
        "About 55 seconds",
        """399 backend tests passed, zero failed, zero skipped. Five frontend Vitest tests. Sixteen frozen domain tables. Zero new tables in Steps 3 through 9.

Live stack we actually ran: PostgreSQL on localhost 5433, API on 127.0.0.1 port 8010, UI on 5173. Health is status ok, service setuhaul. Port 8000 is another local application — we do not use it. The UI rejects a health payload whose service is not setuhaul.

API and model tests use in-memory SQLite. Migration and concurrency tests require real PostgreSQL. Frontend production build passed. alembic check: freeze held.

Hero script: e2e_hero_flow.py. Seed: seed_ops_demo.py for SH-1024.""",
    ),
    (
        "45",
        "Demo — How to run the walkthrough",
        "About 50 seconds",
        """If you are about to click the console, this is the recipe.

Docker compose up for Postgres on 5433. alembic upgrade head. python scripts/seed_ops_demo.py for SH-1024, Jane Rivera, Dallas DC, original 6:30 window and later open slots. uvicorn on port 8010. npm run dev in frontend. Then Demo Scenarios or the six hero messages.

LLM_PROVIDER stays fake unless you intentionally enable OpenRouter.

Reset for a clean confirm: docker compose down dash v, then those steps again. Re-running seed does not undo a confirmed SH-1024. Do not hand-edit appointment rows as a reset procedure.""",
    ),
    (
        "46",
        "Honesty — Known limitations vs out of scope",
        "About 65 seconds",
        """Left column is designed this way, not hidden failures.

No classroom login or JWT. No email, SMS, or push. Escalation is a flag, not a dispatcher SLA. No live travel-time prediction. Proposal TTL is application-side, no expires_at column. No shipment priority or unload-minutes fields. Step 9 assignment is not a hold. dock_id may be null until proposal or allocation. Conversation does not check the driver in at the gate. Direct allocate exists for REST; chat does not use it. Two-row confirm. Vehicle length versus dock max length is not evaluable — no vehicle length field.

Right column is not this architecture: LangChain, LangGraph, Redis, Kafka, OR-Tools, national fleet routing, production IdP, notification platform, human-task inbox, POST schedule confirm, a separate Proposal table.

I would rather list these than pretend we built a production WMS.""",
    ),
    (
        "47",
        "Close — What to remember",
        "About 50 seconds",
        """Six sentences, then stop.

Language versus authority: the model talks, Steps 5 through 7 decide and commit.

Three states: show options, write a proposal, then confirm under locks.

Capacity is scarce: only confirmed and held consume a slot. requested does not.

Revalidate at commit: a proposal can go stale. The loser gets 409, not a silent book.

Frozen facts: 16 tables. Later steps reuse Appointment and ChatMessage metadata.

Step 9 is optional: a ranked schedule is advice. Booking still walks 7, then 5, then 6.

That is SetuHaul.""",
    ),
    (
        "48",
        "Questions — Thank you",
        "About 20 seconds, then Q&A",
        """Thank you. SetuHaul: deterministic appointment coordination. LLM for language, Python for operational authority.

Likely questions, short answers.

“Why not LangChain?” We did not need a graph framework. Allowlisted tools over existing services are the control plane.

“Why no Proposal table?” Schema was frozen at Step 2. Proposals are Appointment requested plus a notes marker.

“Can OpenRouter confirm?” No. Parser flags plus the accept path. The model cannot promote confirm.

“Does Step 9 book?” No. read_only true. No confirm endpoint.

“How do you prevent double-book?” Locks in order, revalidate Step 5, count confirmed and held, one winner, 409 for the loser.

I am happy to walk the Driver Console on SH-1024 live.""",
    ),
]


def _shade(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_run_font(run, name="Calibri", size=12, bold=False, color=INK, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_para(doc, text, *, size=12, bold=False, color=INK, space_after=8, space_before=0, italic=False, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return p


def build() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.95)
        section.right_margin = Inches(0.95)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("SETUHAUL  ·  FDE assignment speaker script  ·  aligned to the 48-slide deck")
        _set_run_font(run, size=9, color=MUTED)

        hdr = section.header
        hdr.is_linked_to_previous = False
        hp = hdr.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run("Speak these paragraphs  ·  do not read the slide bullets verbatim")
        _set_run_font(run, size=9, italic=True, color=GOLD)

    add_para(doc, "SETUHAUL", size=28, bold=True, color=NAVY, space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(
        doc,
        "Presentation speaker script",
        size=18,
        bold=True,
        color=GOLD,
        space_after=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        "Complete spoken explanation for every slide in SetuHaul_FDE_Assignment_Presentation_updated.pptx "
        "(48 slides). Use this document as a teleprompter. Glance at the slide; speak the paragraphs. "
        "Timing is a guide for a ~35–45 minute presentation with a short demo after the hero-flow slides.",
        size=12,
        color=INK,
        space_after=8,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        "How to use this script: Pause on diagram slides and point at the picture while you talk. "
        "If time is cut, keep slides 1, 4, 6, 7, 11, 22, 29, 34, 46, and 47.",
        size=11,
        italic=True,
        color=MUTED,
        space_after=16,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    for num, title, timing, script in SLIDES:
        heading = doc.add_heading(level=1)
        heading.paragraph_format.space_before = Pt(16)
        heading.paragraph_format.space_after = Pt(4)
        run = heading.add_run(f"Slide {num}.  {title}")
        _set_run_font(run, name="Calibri", size=16, bold=True, color=NAVY)

        meta = add_para(doc, f"Time on slide: {timing}", size=11, italic=True, color=GOLD, space_after=8)

        for block in script.strip().split("\n\n"):
            add_para(doc, block.replace("\n", " "), size=12, color=INK, space_after=8)

    doc.save(OUT)
    print(f"Wrote {OUT} ({len(SLIDES)} slides)")


if __name__ == "__main__":
    build()
